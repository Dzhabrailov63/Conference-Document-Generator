# coding: utf8
import pandas as pd
import numpy as np
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Cm


table = pd.read_excel('2024-02-17 Konferentsiia actual.xlsx')

# список с названиями секций
sections_list1 = pd.read_excel('Sections.xlsx')
data2 = sections_list1.values
sections_list = list()
for i in range(0, len(data2)):
    sections_list.append(data2[i])


# список с индексами секций, содержащих подсекции
sect_with_undersect = [0, 2, 3, 18, 21, 26, 36, 40, 76, 100, 103, 104, 106, 107, 116, 122]

# Создание нового документа Word
doc = docx.Document()

# Установка параметров страницы
doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
doc.sections[0].page_width = 10692000
doc.sections[0].page_height = 7559718
section = doc.sections[0]
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
data = np.array([], [])
data = table.values
numberOfRecords = len(data)

# Установка стиля шрифта и размера по умолчанию
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
index_under = int(4)

# Перебор всех секций
for sec in range(0, len(sections_list)):
    # Проверка, содержит ли текущая секция подсекции
    if sec in sect_with_undersect:
        flag_undersect = True
    else:
        flag_undersect = False
    # Добавление заголовка секции
    table = {'Наименование секции': sections_list[sec]}
    head = doc.add_paragraph()
    head.paragraph_format.space_before = 0
    head.paragraph_format.space_after = 0
    run = head.add_run(sections_list[sec])
    run.font.size = Pt(14)
    run.bold = True
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Перебор записей в таблице данных
    for numberRecord in range(numberOfRecords):
        if data[numberRecord][2] == sections_list[sec]:
            # Добавление информации о заседании
            if (sec == 26):
                head = doc.add_paragraph()
                run = head.add_run('Председатель - ' + data[numberRecord][21])
                run.font.size = Pt(12)
                head.paragraph_format.space_before = 0
                head.paragraph_format.space_after = 0
                head.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if flag_undersect:
                table.update({'Подсекция': data[numberRecord][int(index_under)]})
                if (str(data[numberRecord][int(index_under)])) != 'nan':
                    #run = head.add_run('\n' + data[numberRecord][section])
                    print(data[numberRecord][int(index_under)])
                    head = doc.add_paragraph()
                    head.paragraph_format.space_before = 0
                    head.paragraph_format.space_after = 0
                    print(data[numberRecord][2])
                    run = head.add_run(data[numberRecord][int(index_under)])
                    run.font.size = Pt(12)
                    run.bold = True
                    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if str(data[numberRecord][20]) != 'nan':
                    table.update({'Номер заседания подсекции': data[numberRecord][20]})
                    run = head.add_run(' (Заседание ' + str(int(data[numberRecord][20])) + ')')
                run.font.size = Pt(12)
                run.bold = True
            else:
                if (str(data[numberRecord][3]) != 'nan'):
                    table.update({'Номер заседания секции': int(data[numberRecord][3])})

                    head = doc.add_paragraph()
                    head.paragraph_format.space_before = 0
                    head.paragraph_format.space_after = 0
                    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = head.add_run('Заседание № ' + str(int(data[numberRecord][3])))
                    run.font.size = Pt(12)
                    run.bold = True

            # Создание таблицы для данных о заседании
            tableW = doc.add_table(rows=2, cols=0)
            tableW.add_column(width=1058400)
            tableW.add_column(width=6480000)
            tableW.add_column(width=1800000)

            tableW.style = 'First'

            cell = tableW.cell(0, 0)
            cell.text = 'Председатель'
            cell = tableW.cell(0, 1)

            if str(data[numberRecord][25]) != 'nan':
                table.update({'Председатель подсекции': data[numberRecord][25]})
                cell.text = '- ' + data[numberRecord][25]
            elif str(data[numberRecord][21]) != 'nan':
                table.update({'Председатель секции': data[numberRecord][21]})
                cell.text = '- ' + data[numberRecord][21]
            else:
                cell.text = 'Председателя нет..'

            if str(data[numberRecord][26]) != 'nan':
                table.update({'Сопредседатель подсекции': data[numberRecord][26]})
                cell.text += ', сопредседатель - ' + data[numberRecord][26]
            elif str(data[numberRecord][22]) != 'nan':
                table.update({'Сопредседатель секции': data[numberRecord][22]})

                cell.text += ', сопредседатель - ' + data[numberRecord][22]

            cell = tableW.cell(1, 0)
            cell.text = 'Секретарь'
            cell = tableW.cell(1, 1)
            if str(data[numberRecord][27]) != 'nan':
                table.update({'Секретарь подсекции': data[numberRecord][27]})
                cell.text = '- ' + data[numberRecord][27]
            elif str(data[numberRecord][23]) != 'nan':
                table.update({'Секретарь секции': data[numberRecord][23]})
                cell.text = '- ' + data[numberRecord][23]
            else:
                cell.text = 'Секретаря нет'

            if str(data[numberRecord][24]) != 'nan':
                table.update({'Номер телефона секретаря': data[numberRecord][24]})

            cell = tableW.cell(0, 2)
            if str(data[numberRecord][28]) != 'nan':
                table.update({'Дата заседания': data[numberRecord][28]})
                cell.text = 'Заседание '+data[numberRecord][28][9] + ' aпреля'
            else:
                cell.text = 'никогда'

            cell = tableW.cell(1, 2)
            if str(data[numberRecord][29]) != 'nan':
                table.update({'Время': data[numberRecord][29]})
                cell.text = data[numberRecord][29]
            else:
                cell.text = 'нигде и никогда'
            if str(data[numberRecord][31]) != 'nan':
                table.update({'Ссылка на видеоконференцию': data[numberRecord][31]})
                cell.text += ', ' + data[numberRecord][31]
            elif str(data[numberRecord][30]) != 'nan':
                table.update({'Место': data[numberRecord][30]})
                cell.text += ', ' + data[numberRecord][30]
            i = 0
            for speakers in range(32, data.shape[1], 4):
                if str(data[numberRecord][speakers]) != 'nan':
                    i += 1
                else:
                    break

            # Создание таблицы для списка докладчиков
            tableS = doc.add_table(rows=i+1, cols=0)
            tableS.add_column(width=360000)
            tableS.add_column(width=4590000)
            tableS.add_column(width=2520000)
            tableS.add_column(width=1890000)
            # Добавление информации о докладах в таблицу
            i = 0
            tableS.style = 'Second'
            cell = tableS.cell(0, 0)
            cell.text = '№'
            cell = tableS.cell(0, 1)
            cell.text = 'Тема доклада'
            cell = tableS.cell(0, 2)
            cell.text = 'Докладчик'
            cell = tableS.cell(0, 3)
            cell.text = 'Научный руководитель'
            for speakers in range(32, data.shape[1], 4):
                if str(data[numberRecord][speakers]) != 'nan':
                    cell = tableS.cell(i+1, 0)
                    cell.text = str(i+1) + '.'
                    table.update({'Тема доклада [' + str(i) + ']': data[numberRecord][speakers]})
                    cell = tableS.cell(i + 1, 1)

                    if (data[numberRecord][speakers][1] == '.') | (data[numberRecord][speakers][1] == ')'):
                        data[numberRecord][speakers] = data[numberRecord][speakers][2:]
                    if (data[numberRecord][speakers][2] == '.') | (data[numberRecord][speakers][2] == ')'):
                        data[numberRecord][speakers] = data[numberRecord][speakers][3:]
                    if data[numberRecord][speakers][0] == ' ':
                        data[numberRecord][speakers] = data[numberRecord][speakers][1:]
                    cell.text = data[numberRecord][speakers]

                    table.update({'Докладчик[' + str(i) + ']': data[numberRecord][speakers+1]})

                    position = 0
                    while position != -1:
                        position = data[numberRecord][speakers+1].find(".,", position + 1)
                        if position > 0:
                            if data[numberRecord][speakers+1][position + 2] == ' ':
                                data[numberRecord][speakers + 1] = data[numberRecord][speakers+1][:position + 2] +'\n' + data[numberRecord][speakers+1][position + 3:]
                            else:
                                data[numberRecord][speakers+1] = data[numberRecord][speakers+1][:position + 2] + '\n' + data[numberRecord][speakers+1][position + 2:]

                    cell = tableS.cell(i + 1, 2)
                    cell.text = data[numberRecord][speakers+1]
                    table.update({'Научный руководитель[' + str(i) + ']': data[numberRecord][speakers+2]})
                    cell = tableS.cell(i + 1, 3)
                    if str(data[numberRecord][speakers+3]) != 'nan':
                        table.update({'Второй научный руководитель[' + str(i) + ']': data[numberRecord][speakers+3]})
                        cell.text = data[numberRecord][speakers+2] + ',\n' + data[numberRecord][speakers+3]
                    else:
                        cell.text = data[numberRecord][speakers + 2]
                    i+=1
                else:
                    break
            print(table)
    if flag_undersect:
        index_under+=1
        # Сохранение документа

doc.save('test.docx')